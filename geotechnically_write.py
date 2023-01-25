from operator import truediv
import webbrowser
import docx
from docx import Document
import pygame
import datetime
import pandas as pd
import openpyxl




pygame.init()
#button class

clock = pygame.time.Clock()


date = (f'{datetime.datetime.now():%d, %b, %y}')
focus = None
class TextInputBox(pygame.sprite.Sprite):
    def __init__(self, x, y, w, font):
        super().__init__()
        self.color = (255, 255, 255)
        self.pos = (x, y) 
        self.width = w
        self.font = font
        self.active = False
        self.backcolor = None
        self.text = ""
        self.render_text()

    def render_text(self):
        t_surf = self.font.render(self.text, True, self.color, self.backcolor)
        self.image = pygame.Surface((max(self.width, t_surf.get_width()+10), t_surf.get_height()+10), pygame.SRCALPHA)
        if self.backcolor:
            self.image.fill(self.backcolor)
        self.image.blit(t_surf, (5, 5))
        pygame.draw.rect(self.image, self.color, self.image.get_rect().inflate(-2, -2), 2)
        self.rect = self.image.get_rect(topleft = self.pos)

    def update(self, event_list):

        for event in event_list:

            if event.type == pygame.MOUSEBUTTONDOWN and not self.active:
                self.active = self.rect.collidepoint(event.pos)
                self.backcolor = "BLUE"

            if event.type == pygame.MOUSEBUTTONDOWN and self.active:
              self.active = self.rect.collidepoint(event.pos)

                
            if event.type == pygame.KEYDOWN and self.active:
                if event.key == pygame.K_RETURN:
                    self.active = False
                    self.backcolor = "RED"
                elif event.key == pygame.K_BACKSPACE:
                    self.text = self.text[:-1]
                else:
                    self.text += event.unicode
                self.render_text()

font = pygame.font.SysFont(None, 32)
text_input_box_client_name = TextInputBox(50, 50, 400, font)
text_input_box_client_address1 = TextInputBox(50, 115, 400, font)
text_input_box_client_address2 = TextInputBox(50, 180, 400, font)
text_input_box_development_description = TextInputBox(50, 245, 700, font)
text_input_box_nearby_development = TextInputBox(50, 310, 700, font)
text_input_box_north = TextInputBox(50, 375, 400, font)
text_input_box_east = TextInputBox(50, 440, 400, font)
text_input_box_south = TextInputBox(50, 505, 400, font)
text_input_box_west = TextInputBox(50, 570, 400, font)
text_input_box_site_class = TextInputBox(50, 50, 400, font)
text_input_box_ss = TextInputBox(50, 115, 400, font)
text_input_box_s1 = TextInputBox(50, 180, 400, font)
text_input_box_fa = TextInputBox(50, 245, 400, font)
text_input_box_fv = TextInputBox(50, 310, 400, font)
text_input_box_sms = TextInputBox(50, 375, 400, font)
text_input_box_sm1 = TextInputBox(50, 440, 400, font)
text_input_box_sds = TextInputBox(50, 505, 400, font)
text_input_box_sd1 = TextInputBox(50, 570, 400, font)


group2 = pygame.sprite.Group(text_input_box_client_name,text_input_box_client_address1,text_input_box_client_address2,text_input_box_development_description,text_input_box_nearby_development,text_input_box_north,text_input_box_east,text_input_box_south,text_input_box_west)
group1 = pygame.sprite.Group(text_input_box_site_class, text_input_box_ss, text_input_box_s1, text_input_box_fa, text_input_box_fv, text_input_box_sms, text_input_box_sm1, text_input_box_sds, text_input_box_sd1)
# pandas and openpyxl excel file import and setup
excel_file = ('Geo.xlsx')
wb = openpyxl.load_workbook('Geo.xlsx')
point2 = wb['POINT']
project = pd.read_excel(excel_file, sheet_name=0).astype("string")
point = pd.read_excel(excel_file, sheet_name=1).astype("string")
lab_specimen = pd.read_excel(excel_file, sheet_name=2).astype("string")
lithology = pd.read_excel(excel_file, sheet_name=3).astype("string")
sample = pd.read_excel(excel_file, sheet_name=4)
tests = pd.read_excel(excel_file, sheet_name=5)
atterburg = pd.read_excel(excel_file, sheet_name=6).astype("string")
sieve = pd.read_excel(excel_file, sheet_name=7)
wc_density = pd.read_excel(excel_file, sheet_name=8)
attb_readings = pd.read_excel(excel_file, sheet_name=9)
sv_readings = pd.read_excel(excel_file, sheet_name=10)

#gint pandas and openpyxl inputs and conversions
gint_max_depth_drilling = point['HoleDepth'].values.astype(int)
#passed below
gint_max_drilling_depth_corrected_pre = gint_max_depth_drilling.max()
gint_max_drilling_depth_corrected = gint_max_drilling_depth_corrected_pre
# working
gint_drilling_date = str(point2['C2'].value)
#passed below
gint_drilling_date_corrected = gint_drilling_date[:-9]
# working
gint_elevation = point2['E2'].value
# working
gint_hole_size = point2['F2'].value
# working
gint_drilling_contractor = point2['G2'].value
# working
gint_auger_type = point2['H3'].value
#working

gint_client_business_pre = project['Client'].values
gint_client_business = ',\n'.join(gint_client_business_pre)
#working
gint_project_number_pre = project['Number'].values
gint_project_number = ',\n'.join(gint_project_number_pre)
#working
gint_address1_pre = project['Name'].values.astype("string")
gint_address1 = ',\n'.join(gint_address1_pre)
#working
gint_address2_nozip_pre = project['Location'].values
gint_address2_nozip = ',\n'.join(gint_address2_nozip_pre)
#working
gint_plastic_limit_pre_pre = atterburg['Liquid_Limit'].values.astype(float)
gint_plastic_limit = gint_plastic_limit_pre_pre.round(0).astype(int)
gint_plastic_limit_converted = ',\n'.join(gint_plastic_limit.astype(str))

#working
gint_liquid_limit_pre_pre = atterburg['Plastic_Limit'].values.astype(float)
gint_liquid_limit = gint_liquid_limit_pre_pre.round(0).astype(int)
gint_liquid_limit_converted = ',\n'.join(gint_liquid_limit.astype(str))

#working
gint_plasticity_index_pre = gint_plastic_limit - gint_liquid_limit
gint_plasticity_index = ',\n'.join(gint_plasticity_index_pre.astype(str))
#working
gint_depth_atterburg_pre = atterburg['Depth'].values
gint_depth_atterburg = ',\n'.join(gint_depth_atterburg_pre)
#working
gint_atterburg_boring_pre = atterburg['PointID'].values
gint_atterburg_boring = ',\n'.join(gint_atterburg_boring_pre)
#working
gint_boring_names_pre = point['PointID'].astype(str).values.tolist()
gint_boring_names = ', '.join(gint_boring_names_pre)

#working
gint_number_of_borings = len(point['PointID'].values)
#working

if gint_plasticity_index_pre == 0.0:
  gint_plasticity_check = 'non'
if gint_plasticity_index_pre > 0.0 and gint_plasticity_index_pre < 7.0:
  gint_plasticity_check = 'low'
if gint_plasticity_index_pre >= 7.0 and gint_plasticity_index_pre <= 17.0:
  gint_plasticity_check = 'medium'
if gint_plasticity_index_pre >= 17.0:
  gint_plasticity_check = 'high'
  expansive_hazard = True
#working  
# value list to be modified and referred to for document creation

client_name = text_input_box_client_name
client_address1 = text_input_box_client_address1.text
client_address2 = text_input_box_client_address2.text
project_number = gint_project_number
project_address1 = gint_address1
project_address2 = gint_address2_nozip
geomorphic_region = ''
demolition = False
development_description = text_input_box_development_description.text
north = text_input_box_north.text
east = text_input_box_east.text
south = text_input_box_south.text
west = text_input_box_west.text
topography = ''
local_development = text_input_box_nearby_development.text
local_development_capitalized = text_input_box_nearby_development.text.title
expansive_hazard = False
liquifaction_hazard = False
rupture_hazard = False
landslide_hazard = False
gint = False
prefix = ''

class Button():
	def __init__(self, x, y, image, scale):
		width = image.get_width()
		height = image.get_height()
		self.image = pygame.transform.scale(image, (int(width * scale), int(height * scale)))
		self.rect = self.image.get_rect()
		self.rect.topleft = (x, y)
		self.clicked = False

	def draw(self, surface):
		action = False
		#get mouse position
		pos = pygame.mouse.get_pos()

		#check mouseover and clicked conditions
		if self.rect.collidepoint(pos):
			if pygame.mouse.get_pressed()[0] == 1 and self.clicked == False:
				self.clicked = True
				action = True

		if pygame.mouse.get_pressed()[0] == 0:
			self.clicked = False

		#draw button on screen
		surface.blit(self.image, (self.rect.x, self.rect.y))

		return action

#create game window
SCREEN_WIDTH = 1200
SCREEN_HEIGHT = 720

screen = pygame.display.set_mode((SCREEN_WIDTH, SCREEN_HEIGHT))
pygame.display.set_caption("Main Menu")

#define fonts
font = pygame.font.SysFont("arialblack", 40)

#define colours
TEXT_COL = (255, 255, 255)

#load button images
resume_img = pygame.image.load("light_blue.bmp").convert_alpha()
clicked_resume_img = pygame.image.load("dark_blue.bmp")
title_font = pygame.font.Font(None,50)
text_font = pygame.font.Font(None,40)
text_font24 = pygame.font.Font(None,24)
title_text = title_font.render('Geotechnically Write V.1', False, 'Blue')
main_menu_text1 = text_font.render('New Geotechnical Report', False, 'Blue')
main_menu_text2 = text_font.render('Quit', False, 'Blue')
main_menu_text3 = text_font.render('New Geotechnical Report (Learning Mode)', False, 'Blue')
textquit = text_font.render("Quit", False, 'Blue')
textback = text_font.render("Back", False, 'Blue')
textnext = text_font.render("Next", False, 'Blue')
text2 = text_font.render("Today's  Date", False, 'Blue')
text3 = text_font.render('Client Name, Bob Dole', False, 'Blue')
text4 = text_font.render('Clients Mailing Address Line 1, 123 Fake St', False, 'Blue')
text5 = text_font.render('Client Mailing Address Line 2, Ice Cream Island, Alaska 12345', False, 'Blue')
text6 = text_font.render('Project Number', False, 'Blue')
text7 = text_font.render('Project Address Line 1, i.e 123 Fake St', False, 'Blue')
text8 = text_font.render('Project Address Line 2, i.e Ice Cream Island, Alaska 12345', False, 'Blue')
text9 = text_font.render('Choose Geomorphic Region', False, 'Blue')
text10 = text_font.render('Are Any Structures on Site Getting Demolished?', False, 'Blue')
text11 = text_font.render('Development Discription, i.e " three-story multifamily residence"', False, 'Blue')
text12 = text_font.render('What is North of the Subject Property', False, 'Blue')
text13 = text_font.render('What is East of the Subject Property', False, 'Blue')
text14 = text_font.render('What is South of the Subject Property', False, 'Blue')
text15 = text_font.render('What is West of the Subject Property', False, 'Blue')
text16 = text_font.render('How steep is General Topography of the Site', False, 'Blue')
text16f = text_font.render('Flat', False, 'Blue')
text16g = text_font.render('Gentle', False, 'Blue')
text16s = text_font.render('Steep', False, 'Blue')
text17 = text_font.render('Type of Nearby Development, i.e residential or agricultural land', False, 'Blue')
text18 = text_font.render('First Local Major Fault', False, 'Blue')
text19 = text_font.render('Second Local Major Fault', False, 'Blue')
text20 = text_font.render('Third Local Major Fault', False, 'Blue')
text21 = text_font.render('Distance to Fault', False, 'Blue')
text22 = text_font.render('Distance to Fault', False, 'Blue')
text23 = text_font.render('Distance to Fault', False, 'Blue')
text24 = text_font.render('State or County Identified Landslide Hazard Zone?', False, 'Blue')
# text25 = text_font.render('State or County Identified Expansive Soil Hazard Zone?', False, 'Blue')
text26 = text_font.render('State or County Identified Liquifation Hazard Zone?', False, 'Blue')
text27 = text_font.render('State or County Identified Fault Rupture Hazard Zone?', False, 'Blue')
text28 = text_font.render('Click Button to Prepare Geotechnical Engineering Report', False, 'Blue')
text29 = text_font.render('Program Will Close and Report Will be in Folder After Closing', False, 'Blue')
textyes = text_font.render('Yes', False, 'Blue')
textno = text_font.render('No', False, 'Blue')
textfinish = text_font.render('Finish', False, 'Blue')
textgv = text_font24.render('Great Valley', False, 'Blue')
textmp = text_font24.render('Modoc Plateau', False, 'Blue')
textsn = text_font24.render('Sierra Nevada', False, 'Blue')
textkm = text_font24.render('Klamath Mountains', False, 'Blue')
texttr = text_font24.render('Transverse', False, 'Blue')
textcr = text_font24.render('Cascade Range', False, 'Blue')
textcoast = text_font24.render('Coast Ranges', False, 'Blue')
textpr = text_font24.render('Peninsular Ranges', False, 'Blue')
textbr = text_font24.render('Basin and Ranges', False, 'Blue')
textmd = text_font24.render('Mojave Desert', False, 'Blue')
textcd = text_font24.render('Colorado Desert', False, 'Blue')
textflat = text_font24.render('Flat', False, 'Blue')
textgentle = text_font24.render('Gentle', False, 'Blue')
textsteep = text_font24.render('Steep', False, 'Blue')
textsite_class = text_font.render('Enter Site Class', False, 'Blue')
textss = text_font.render('Enter Ss Value', False, 'Blue')
texts1 = text_font.render('Enter S1 Value', False, 'Blue')
textFa = text_font.render('Enter Fa Value', False, 'Blue')
textFv = text_font.render('Enter Fv Value', False, 'Blue')
textSms = text_font.render('Enter Sms Value', False, 'Blue')
textsm1 = text_font.render('Enter Sm1 Value', False, 'Blue')
textSds = text_font.render('Enter Sds Value', False, 'Blue')
textSd1 = text_font.render('Enter Sd1 Value', False, 'Blue')

#create button instances

back_button = Button(25, 645, resume_img, .1)
start_button = Button(525, 190, resume_img, .1)
learning_start_button = Button(525, 315, resume_img, .1)
forward_button = Button(1075, 645, resume_img, .1)
quit_button = Button(1075, 25, resume_img, .1)
demolitionyes_button = Button(375, 150, resume_img, .1)
demolitionyes_button_clicked = Button(375, 150, clicked_resume_img, .1)
demolitionno_button = Button(800, 150, resume_img, .1)
demolitionno_button_clicked = Button(800, 150, clicked_resume_img, .1)
great_valley_button = Button(100, 135, resume_img, .1)
great_valley_button_clicked = Button(100, 135, clicked_resume_img, .1)
modoc_button = Button(100, 185, resume_img, .1)
modoc_button_clicked = Button(100, 185, clicked_resume_img, .1)
sierra_button = Button(100, 235, resume_img, .1)
sierra_button_clicked = Button(100, 235, clicked_resume_img, .1)
klamath_button = Button(400, 135, resume_img, .1)
klamath_button_clicked = Button(400, 135, clicked_resume_img, .1)
transverse_button = Button(400, 185, resume_img, .1)
transverse_button_clicked = Button(400, 185, clicked_resume_img, .1)
cascade_button = Button(400, 235, resume_img, .1)
cascade_button_clicked = Button(400, 235, clicked_resume_img, .1)
coast_button = Button(700, 135, resume_img, .1)
coast_button_clicked = Button(700, 135, clicked_resume_img, .1)
peninsular_button = Button(700, 185, resume_img, .1)
peninsular_button_clicked = Button(700, 185, clicked_resume_img, .1)
basin_range_button = Button(700, 235, resume_img, .1)
basin_range_button_clicked = Button(700, 235, clicked_resume_img, .1)
mojave_button = Button(1000, 135, resume_img, .1)
mojave_button_clicked = Button(1000, 135, clicked_resume_img, .1)
colorado_button = Button(1000, 185, resume_img, .1)
colorado_button_clicked = Button(1000, 185, clicked_resume_img, .1)
flat_button = Button(300, 50, resume_img, .1)
flat_button_clicked = Button(300, 50, clicked_resume_img, .1)
gentle_button = Button(600, 50, resume_img, .1)
gentle_button_clicked = Button(600, 50, clicked_resume_img, .1)
steep_button = Button(900, 50, resume_img, .1)
steep_button_clicked = Button(900, 50, clicked_resume_img, .1)
landslideyes_button = Button(365, 340, resume_img, .1)
landslideyes_button_clicked = Button(365, 340, clicked_resume_img, .1)
landslideno_button = Button(660, 340, resume_img, .1)
landslideno_button_clicked = Button(660, 340, clicked_resume_img, .1)
# expansiveyes_button = Button(365, 405, resume_img, .1)
# expansiveyes_button_clicked = Button(365, 405, clicked_resume_img, .1)
# expansiveno_button = Button(660, 405, resume_img, .1)
# expansiveno_button_clicked = Button(660, 405, clicked_resume_img, .1)
liquifactionyes_button = Button(365, 440, resume_img, .1)
liquifactionyes_button_clicked = Button(365, 440, clicked_resume_img, .1)
liquifactionno_button = Button(660, 440, resume_img, .1)
liquifactionno_button_clicked = Button(660, 440, clicked_resume_img, .1)
faultyes_button = Button(365, 540, resume_img, .1)
faultyes_button_clicked = Button(365, 540, clicked_resume_img, .1)
faultno_button = Button(660, 540, resume_img, .1)
faultno_button_clicked = Button(660, 540, clicked_resume_img, .1)
gintyes_button = Button(550, 200, resume_img, .1)
gintyes_button_clicked = Button(400, 200, clicked_resume_img, .1)
gintno_button = Button(600, 200, resume_img, .1)

doc = docx.Document('main_report.docx')

#learning Mode Variable
learning_mode = False

#game loop and webpage loading reset variables
run = True
screen3 = False
screen4 = False
screen5 = False
screen6 = False
screen7 = False

#gint variables
gintno = False
gintyes = False

#topographic variables
steep = False
gentle = False
flat = False

#geomorphic variables
colorado = False
mojave = False
basin = False
peninsular = False
coast = False
cascade = False
transverse = False
klamath = False
sierra = False
modoc = False
great_valley = False


#seismic variables
site_class_description = ''
site_class = text_input_box_site_class.text
if text_input_box_site_class.text == 'a' or 'A':
  site_class_description = 'hard rock'
elif text_input_box_site_class.text == 'b' or 'B':
  site_class_description = 'rock'
elif text_input_box_site_class.text == 'c' or 'C':
  site_class_description = 'very dense soil and hard rock'
elif text_input_box_site_class.text == 'd' or 'D':
  site_class_description = 'stiff soil'
elif text_input_box_site_class.text == 'e' or 'E':
  site_class_description = 'soft clay soil'
else:
  pass


ss = text_input_box_ss.text
s1 = text_input_box_s1.text
fa = text_input_box_fa.text
fv = text_input_box_fv.text
sms = text_input_box_sms.text
sm1 = text_input_box_sm1.text
sds = text_input_box_sds.text
sd1 = text_input_box_sd1.text

seismic_analysis_required = False


project_start = False
menu_state = "2"

while run:
  clock.tick(8)
  screen.fill((0, 0, 0))
  
  #check if program is paused
  if project_start == True:
    
    #check menu state
    if menu_state == "2":
      #draw pause screen buttons
      if back_button.draw(screen):
        project_start = False
      
      if forward_button.draw(screen):
        menu_state = "3"        
      if quit_button.draw(screen):
        run = False
      
      event_list = pygame.event.get()
      group2.update(event_list)
      group2.draw(screen)


      screen.blit(text3,(50,20))
      screen.blit(text4,(50,85))
      screen.blit(text5,(50,150))
      screen.blit(text11,(50,215))
      screen.blit(text17,(50,280))    
      screen.blit(text12,(50,345))
      screen.blit(text13,(50,410))
      screen.blit(text14,(50,475))
      screen.blit(text15,(50,540))
      screen.blit(textquit,(1100,35))
      screen.blit(textback,(50,655))
      screen.blit(textnext,(1100,655))
      if learning_mode == True:
        if screen4 == False:
          webbrowser.open('https://earth.google.com/', new = 2)
          screen4 = True


    #check if the menu is open
    if menu_state == "3":
      #draw the different  buttons
      if flat_button.draw(screen):
        topography = ('flat')
      if gentle_button.draw(screen):
        topography = ('gentle')
      if steep_button.draw(screen):
        topography = ('steep')
      if great_valley_button.draw(screen):
        colorado = False
        mojave = False
        basin = False
        peninsular = False
        coast = False
        cascade = False
        transverse = False
        klamath = False
        sierra = False
        modoc = False
        great_valley = True
      if modoc_button.draw(screen):
        colorado = False
        mojave = False
        basin = False
        peninsular = False
        coast = False
        cascade = False
        transverse = False
        klamath = False
        sierra = False
        modoc = True
        great_valley = False
      if sierra_button.draw(screen):
        colorado = False
        mojave = False
        basin = False
        peninsular = False
        coast = False
        cascade = False
        transverse = False
        klamath = False
        sierra = True
        modoc = False
        great_valley = False
      if klamath_button.draw(screen):
        colorado = False
        mojave = False
        basin = False
        peninsular = False
        coast = False
        cascade = False
        transverse = False
        klamath = True
        sierra = False
        modoc = False
        great_valley = False
      if transverse_button.draw(screen):
        colorado = False
        mojave = False
        basin = False
        peninsular = False
        coast = False
        cascade = False
        transverse = True
        klamath = False
        sierra = False
        modoc = False
        great_valley = False
      if cascade_button.draw(screen):
        colorado = False
        mojave = False
        basin = False
        peninsular = False
        coast = False
        cascade = True
        transverse = False
        klamath = False
        sierra = False
        modoc = False
        great_valley = False
      if peninsular_button.draw(screen):
        colorado = False
        mojave = False
        basin = False
        peninsular = True
        coast = False
        cascade = False
        transverse = False
        klamath = False
        sierra = False
        modoc = False
        great_valley = False
      if coast_button.draw(screen):
        colorado = False
        mojave = False
        basin = False
        peninsular = False
        coast = True
        cascade = False
        transverse = False
        klamath = False
        sierra = False
        modoc = False
        great_valley = False
      if basin_range_button.draw(screen):
        colorado = False
        mojave = False
        basin = True
        peninsular = False
        coast = False
        cascade = False
        transverse = False
        klamath = False
        sierra = False
        modoc = False
        great_valley = False
      if mojave_button.draw(screen):
        colorado = False
        mojave = True
        basin = False
        peninsular = False
        coast = False
        cascade = False
        transverse = False
        klamath = False
        sierra = False
        modoc = False
        great_valley = False
      if colorado_button.draw(screen):
        colorado = True
        mojave = False
        basin = False
        peninsular = False
        coast = False
        cascade = False
        transverse = False
        klamath = False
        sierra = False
        modoc = False
        great_valley = False

      if topography == 'gentle':
        gentle_button_clicked.draw(screen)
      if topography == 'flat':
        flat_button_clicked.draw(screen)
      if topography == 'steep':
        steep_button_clicked.draw(screen)

      if colorado == True:
        colorado_button_clicked.draw(screen)

      if mojave == True:
        mojave_button_clicked.draw(screen)

      if basin == True:
        basin_range_button_clicked.draw(screen)

      if peninsular == True:
        peninsular_button_clicked.draw(screen)

      if coast == True:
        coast_button_clicked.draw(screen)

      if cascade == True:
        cascade_button_clicked.draw(screen)

      if transverse == True:
        transverse_button_clicked.draw(screen)

      if klamath == True:
        klamath_button_clicked.draw(screen)

      if sierra == True:
        sierra_button_clicked.draw(screen)

      if modoc == True:
        modoc_button_clicked.draw(screen)

      if great_valley == True:
        great_valley_button_clicked.draw(screen)

      if forward_button.draw(screen):
        menu_state = "4" 

      if landslideyes_button.draw(screen):
        landslide_hazard = True
      if landslideno_button.draw(screen):
        landslide_hazard = False
      if landslide_hazard == False:    
        landslideno_button_clicked.draw(screen)
      if landslide_hazard == True:
        landslideyes_button_clicked.draw(screen)

      if liquifactionyes_button.draw(screen):
        liquifaction_hazard = True
      if liquifactionno_button.draw(screen):
        liquifaction_hazard = False
      if liquifaction_hazard == True:
        liquifactionyes_button_clicked.draw(screen)
      if liquifaction_hazard == False:
        liquifactionno_button_clicked.draw(screen)

      if faultyes_button.draw(screen):
        rupture_hazard = True
      if faultno_button.draw(screen):
        rupture_hazard = False
      if rupture_hazard == True:
        faultyes_button_clicked.draw(screen)
      if rupture_hazard == False:
        faultno_button_clicked.draw(screen)


      event_list = pygame.event.get()
      screen.blit(textflat,(335,65))
      screen.blit(textgentle,(625,65))
      screen.blit(textsteep,(935,65))  
      screen.blit(textgv,(100,150))
      screen.blit(textmp,(100,200))
      screen.blit(textkm,(400,150))
      screen.blit(textsn,(100,250))
      screen.blit(texttr,(400,200))
      screen.blit(textcr,(400,250))
      screen.blit(textcoast,(700,150))
      screen.blit(textpr,(700,200))
      screen.blit(textbr,(700,250))
      screen.blit(textmd,(1000,150))
      screen.blit(textcd,(1000,200))
      screen.blit(text9,(425,100))
      screen.blit(text16,(325,15))
      screen.blit(text24,(225,300))
      screen.blit(text26,(225,400))
      screen.blit(text27,(225,500))
      screen.blit(textyes,(400,350))
      screen.blit(textyes,(400,450))
      screen.blit(textyes,(400,550))
      screen.blit(textno,(700,350))
      screen.blit(textno,(700,450))
      screen.blit(textno,(700,550))
      if learning_mode == True:
        if screen6 == False:
          webbrowser.open('https://maps.conservation.ca.gov/cgs/EQZApp/app/', new = 2)
          screen6 = True

      
      if back_button.draw(screen):
        menu_state = "2"
      if quit_button.draw(screen):
        run = False
       
      screen.blit(textquit,(1100,35))
      screen.blit(textback,(50,655))
      screen.blit(textnext,(1100,655))
      
    if menu_state == "4":
      event_list = pygame.event.get()
      group1.update(event_list)
      group1.draw(screen)
      if forward_button.draw(screen):
        menu_state = "5" 
      if back_button.draw(screen):
        menu_state = "3"
      if quit_button.draw(screen):
        run = False
      if learning_mode == True:
        if screen7 == False:
          webbrowser.open('https://www.seismicmaps.org/', new = 2)
          screen7 = True

      screen.blit(textsite_class,(50,25))
      screen.blit(textss,(50,90))
      screen.blit(texts1,(50,155))
      screen.blit(textFa,(50,220))
      screen.blit(textFv,(50,285))
      screen.blit(textSms,(50,350))
      screen.blit(textsm1,(50,415))
      screen.blit(textSds,(50,480))
      screen.blit(textSd1,(50,545))
      screen.blit(textquit,(1100,35))
      screen.blit(textnext,(1100,655))
      screen.blit(textback,(50,655))

    if menu_state == "5":
      if back_button.draw(screen):
        menu_state = "4"
      if quit_button.draw(screen):
        run = False
      if gintyes_button.draw(screen):
        gintyes = True
        if gintyes == True:
          #start writing to doc     
          # logic before writing
          # 
          
          s1_float = float(text_input_box_s1.text)
          if site_class == 'd' or 'D' and s1_float >= .2:
            seismic_analysis_required = True
          else:
            seismic_analysis_required = False  





          development_description_1 = doc.paragraphs[12]
          development_description_1.add_run(f'{text_input_box_development_description.text.title()} Development')
          address1_1 = doc.paragraphs[13]
          address1_1.add_run(gint_address1)
          address2_1 = doc.paragraphs[14]
          address2_1.add_run(gint_address2_nozip)
          date_1 = doc.paragraphs[15]
          date_1.add_run(date)
          client_business_1 = doc.paragraphs[23]
          client_business_1.add_run(gint_client_business)
          client_address_1 = doc.paragraphs[24]
          client_address_1.add_run(text_input_box_client_address1.text)
          client_address_2 = doc.paragraphs[25]
          client_address_2.add_run(text_input_box_client_address2.text)
          project_number_1 = doc.paragraphs[35]
          project_number_1.add_run(gint_project_number)
          date_2 = doc.paragraphs[38]
          date_2.add_run(date)
          client_business_2 = doc.paragraphs[40]
          client_business_2.add_run(gint_client_business)
          client_address_3 = doc.paragraphs[41]
          client_address_3.add_run(text_input_box_client_address1.text)
          client_address2_3 = doc.paragraphs[42]
          client_address2_3.add_run(text_input_box_client_address2.text)
          client_name_3 = doc.paragraphs[43]
          client_name_3.add_run(text_input_box_client_name.text)
          development_description_4 = doc.paragraphs[46]
          development_description_4.add_run(text_input_box_development_description.text.title() + ' Development')
          address1_2 = doc.paragraphs[47]
          address1_2.add_run(gint_address1)
          address2_2 = doc.paragraphs[48]
          address2_2.add_run(gint_address2_nozip)
          project_number_3 = doc.paragraphs[49]
          project_number_3.add_run(gint_project_number)
          client_name_2 = doc.paragraphs[51]
          client_name_2.add_run(f' {text_input_box_client_name.text},')
          project_address1_1 = doc.paragraphs[53]
          project_address1_1.add_run(f' has prepared a Geotechnical Engineering Study for the proposed {text_input_box_development_description.text} development at the property located at {gint_address1} in {project_address2}. It is our understanding that the proposed development consists of the construction of a {text_input_box_development_description.text}.')
          client_business_3 = doc.paragraphs[55]
          client_business_3.add_run(f'Should you or members of the design team have questions or need additional information, please contact the undersigned at (925) 433-0450 or by e-mail at . We greatly appreciate the opportunity to be of service to {gint_client_business}, and to be involved in the design of this project.')
          project_address1_2 = doc.paragraphs[103]
          project_address1_2.add_run(f'The proposed improvement project is located at {gint_address1} in {gint_address2_nozip} as shown on Figure 1, Site Vicinity Map. The project site is bordered by {text_input_box_north.text} to the north, {text_input_box_east.text} to the east, and {text_input_box_south.text} to the south, and {text_input_box_west.text} to the west. The topography of the site is relatively {topography}, with approximate elevations of {gint_elevation} feet above sea level. The average geographical coordinates used in our engineering analyses are '''' degrees north latitude and '''' degrees west longitude.')
          development_description_2 = doc.paragraphs[108]
          development_description_2.add_run(f' It is our understanding that the proposed development consists of the construction of a {text_input_box_development_description.text} development, as shown on Figure 2, Site Development Plan. In addition to construction of the {text_input_box_development_description.text} development, there will be various associated site improvements such as grading, landscaping, paving, and utilities.')
          drilling_info_1 = doc.paragraphs[113]
          drilling_info_1.add_run(f'Our field exploration program consisted of drilling {gint_number_of_borings} soil borings as shown on Figure 3, Site Map and Boring Locations.')
          drilling_info_1.add_run(' A Geo-Eng Staff Engineer visually classified the materials encountered in the borings according to the Unified Soil Classification System as the borings were advanced. Relatively undisturbed soil samples were recovered at selected intervals using a three-inch outside diameter Modified California split spoon sampler containing six-inch long brass liners. A two-inch outside diameter Standard Penetration Test (SPT) sampler was also used to obtain SPT blow counts and obtain disturbed soil samples. The samplers were driven by using a 140-pound safety hammer with an approximate 30-inch fall utilizing N-rods as necessary. Resistance to penetration was recorded as the number of hammer blows required to drive the sampler the final foot of an 18-inch drive. All the blow counts recorded using Modified California split spoon samplers in the field were converted to equivalent SPT blow counts using appropriate modification factors suggested by Burmister (1948), i.e., a factor of 0.65 assuming an inner diameter of 2.5 inches. Therefore, all blow counts shown on the final boring logs are either directly measured (SPT sampler) or equivalent SPT (MC sampler) blow counts. Bulk samples were obtained in the upper few feet of the borings from the auger cuttings as needed. ')
          drilling_info_1.add_run(f' The {gint_number_of_borings} borings were drilled at the site on {gint_drilling_date_corrected}, by {gint_drilling_contractor}, using a truck mounted drill rig equipped with {gint_hole_size} diameter {gint_auger_type}, to a maximum depth of {gint_max_drilling_depth_corrected} feet below existing ground surface.')
          
          drilling_info_1.add_run(f'During our subsurface exploration program, we investigated the subsurface soils and evaluated soil conditions to a maximum depth of {gint_max_drilling_depth_corrected} feet in the borings performed for this study. From the ground surface to the maximum depth explored, the soils underlying the project site consist primarily of a layer of (eventual soil density and hardness check if time allows) below ground surface.')
          # atterburg_1 = doc.paragraphs[139]class
          # atterburg_1.add_run(f'Soil sample of the near surface fine grained material from Boring {gint_atterburg_boring} at {gint_depth_atterburg} feet below ground surface was tested for Atterberg Limits, with measured Liquid Limit (LL) of {gint_liquid_limit}, Plastic Limit (PL) of {gint_plastic_limit}, and corresponding Plasticity Index (PI) of 16. Based on these test results the near surface soil would be considered to be of {gint_plasticity_check} plasticity and have a {gint_plasticity_check} expansion potential.')
          liquifaction_1 = doc.paragraphs[159]
          liquifaction_1.add_run(f'The soils encountered in the subsurface investigation included layers of (soil lithology and density check if timing allows). These soils are expected to be generally (compare if coarse_grain = True) susceptible to liquefaction due to their (fine/coarse check if timing allows) content and relatively (density_high = True check if timing allows) density. Additionally, check if groundwater is encountered. ')
          liquifaction_2 = doc.paragraphs[159]
          liquifaction_2.add_run(f' Therefore, the potential for liquefaction of the site subsurface soils is judged to be (logic check if timing allows).')
          dynamic_compaction_1 = doc.paragraphs[161]
          dynamic_compaction_1.add_run('Dynamic compaction is a phenomenon where loose, relatively clean, near-surface sandy soil located above the water table is densified from vibratory loading, typically from strong seismic shaking or vibratory equipment. The site soils generally consist of (soil lithology and hardness check if timing allows). Therefore, in our opinion, dynamic settlement and/or any potential effect of dynamic settlement on the proposed construction (logic check if timing allows) expected to be (logic check if timing allows).')
          # rupture_hazard_2 = doc.paragraphs[173]
          # rupture_hazard_2.add_run(f'{rupture_hazard_high_low}.')
          # can change with some work to make recommendations.


          atterburg_2 = doc.paragraphs[174]
          atterburg_2.add_run(f'{gint_plasticity_check} expansive (gint_coarseness_check if time permits)-grained soils were encountered in the upper five feet during our subsurface exploration. The results of the laboratory testing performed on a representative sample of the most expansive near-surface soils indicated a measured Plasticity Index of {gint_plasticity_index}, indicative of a {gint_plasticity_check} plasticity and {gint_plasticity_check} expansion potential. "Discuss recommendations for {gint_plasticity_check} is recommended for this site."')
          seismic_coeffecients_1 = doc.paragraphs[173]
          seismic_coeffecients_1.add_run(f'The subject site is located within a seismically active region and should be designed to account for earthquake ground motions as described in this report. Based on the subsurface conditions encountered and our evaluation of the geology of the site, Site Class “{text_input_box_site_class}”, representative of {site_class_description} averaged over the uppermost 100 feet of the subsurface profile would be appropriate for this site.')


          rupture_hazard_1 = doc.paragraphs[131]
          if rupture_hazard == True:
            rupture_hazard_1.add_run('The site IS currently within a designated Earthquake Fault Zone as defined by the State (Hart and Bryant, 1997) or any local zone. Based on our evaluation, a fault investigation study will be necassary')
          else:
            rupture_hazard_1.add_run('The site is NOT currently within a designated Earthquake Fault Zone as defined by the State (Hart and Bryant, 1997) or any local zone. Based on our evaluation, the potential for fault ground rupture or creep at the site is NIL')
          liquifaction_hazard_1 = doc.paragraphs[166]
          if liquifaction_hazard == True:
            liquifaction_hazard_1.add_run('The site IS mapped by by the CGS in a geologic hazard zone requiring liquifaction investigation. <<Enter Description of Soil and Groundwater levels>>')
          if liquifaction_hazard == False:
            liquifaction_hazard_1.add_run('The site is NOT mapped by the CGS in a geologic hazard zone requiring liquifaction investigation. Therefor the potential for liquefaction of the subsurface soil is judged to be low.')
          geomorphic_region1 = doc.paragraphs[123]
          if coast == True:            
            geomorphic_region1.add_run('The site is located within the Coast Ranges geomorphic province of California. The Coast Ranges geomorphic province consists of numerous small to moderate linear mountain ranges trending north to south and northwest to southeast. The Coast Ranges lies between the Pacific Ocean to the west and the Great Valley Geomorphic Province to the east. This province is approximately 400 miles long and extends from the Klamath Mountains in the north to the Santa Ynez River within Santa Barbara County in the south. It generally consists of marine sedimentary rocks and volcanic rocks. The province is characterized by northwest-trending faults and folds, as well as erosion and deposition within the broad transform boundary between the North American and Pacific plates. Translational motion along the plate boundary occurs across a distributed zone of right-lateral shear expressed as a nearly 50-mile-wide zone of northwest-trending, near-vertical active strike-slip faults. This motion occurs primarily along the active San Andreas, Hayward, Calaveras and San Gregorio faults. Site specific geologic units and nearby faults are shown on, Figure 4-Site Vicinity Geologic Map.')
          if transverse == True:
            geomorphic_region1.add_run('The site is in the Transverse Range geomorphic province of California which extends from the southern end of the Coast Range near Santa Maria to the west to the southern end of the Sierra Nevada Mountains to the northwest of the site.  The Transverse Range is an east-west trending geomorphic province that is controlled by the San Andreas Fault which forms the major structural feature of the northern extent of the province.  The northern part of the range is dominated by steeply side predominantly Mesozoic Granite of the San Gabriel Mountains.  The southern part of the Range is dominate by the ranges of the Los Angeles Basin.  The Los Angeles basin is divided into four structural blocks with northwest trending anticlines and synclines.  The subject site is in the Chino basin which is located on the far eastern extent of the Northeastern Block. Site specific geologic units and nearby faults are shown on, Figure 4-Site Vicinity Geologic Map.')
          if great_valley == True:
            geomorphic_region1.add_run('The site is located within the Great Valley physiographic province of California. The Great Valley physiographic province is a large depositional valley consisting of sediments sourced from the Coast Ranges to the west and the Sierra Nevada to the east. This province is approximately 400 miles long and extends from the Klamath Mountains and the Cascade Range in the north to the Transverse Ranges to the south. Sediments were deposited in this valley primarily during the late Mesozoic Era when the valley comprised an ancient seaway and later in the Cenozoic from river deposition. Underlying these sediments are Franciscan Assemblage rocks from the subduction of the Farallon plate under the North American plate. Site specific geologic units and nearby faults are shown on, Figure 4-Site Vicinity Geologic Map.')
          if modoc == True:
            geomorphic_region1.add_run('The site is located within the Modoc Plateau geomorphic province of California. The province is a volcanic table land of high elevation consisting of thick accumaltion of lava flows and tuff bes along with many small volcanic cones. Occasional lakes, marshesm and sluggeshly flowing streams meander across the plateau. The plateau is cut vby many north-south faults. The proving is bound indefinitely by the Cascade Range on the west and the Basin and Range on the east and south. Site specific geologic units and nearby faults are shown on, Figure 4-Site Vicinity Geologic Map.')
          if sierra == True:
            geomorphic_region1.add_run('The site is located within the Sierra Nevada geomorphic province of California. The Sierra Nevada is a tilted fault block nearly 400 miles long. Its east face is a high, rugged, multiple scarp, contrasting with the gentle western slope. Their upper courses, especially in massive granites of the higher Sierra, are modified by glacier sculpturing, forming such scenic features as Yosimite Valley. The high crest culminates in Mt. Whitney at an elevation of 14,495 feet above sea level near the eastern scarp. The metamorphic bedrock contains gold bearing veins in the northwest trending Mother Lode. The northern Sierra boundary is marked where bedrock dissappears under the cenozoic volcanic cover of the Cascade range, the Great Valley province on the west, and the Basin and Range province to its east. Site specific geologic units and nearby faults are shown on, Figure 4-Site Vicinity Geologic Map.')  
          if klamath == True:
            geomorphic_region1.add_run('The site is located within the Klamath Mountains geomorphic province of California. The Klamath Mountains have rugged topography with prominent peaks reaching 6,000-8,000 feet above sea level. In  the wester Klamath Mountains an irregular frainage is incised into an uplifted plateau called the Klamath peneplain. The uplift has left successive benches with gold-bearing gravels on the sides of the canyons. The Klamath River follows a circuitous course from the Cascade Range through the Klamath Mountains. The province is bordered to the south and west by the Coast Range geomorphic province, and the Cascade Range to its east. Site specific geologic units and nearby faults are shown on, Figure 4-Site Vicinity Geologic Map.')
          if basin == True:
            geomorphic_region1.add_run('The site is located within the Basin and Range geomorphic province of California. The Basin and Range is the westernmost part of the Great Basin. The province is characterized by interior drainage with lakes and playas, and the typical horst and graben structure (subparrallelm fault-bounded ranges separated by downdropped basins). Death Valley, the lowest area in the United States (280 feet below sea level at Badwater), is one of the grabens. Another graben, Owens Valley, lies between the bold eastern fault scarp of the Sierra Nevada and Inyo Mountains. The province is bordered to the west by the Sierra Nevada province and the Mojave Desert province to its south. Site specific geologic units and nearby faults are shown on, Figure 4-Site Vicinity Geologic Map.')
          if mojave == True:
            geomorphic_region1.add_run('The site is located with the Mojave Desert geomorphic province of California. The Mojave is a broad interior region of isolated mountain ranges seperated by expanses of desert plains. It has in interior enclosed drainage and many playas. There are two important fault trends that control topography- a prominent NW-SE trend and a secondary east-west trend (apparent alighment with Transverse Ranges is significant). The Mojave province is wedged in a sharp angle between the Garlock Fault (southern boundary Sierra Nevada) and the Sand Andreas Faultm where it bends east from its northwest trend. The northern boundary of the Mojave is separated from the prominent Basin and Range by the easter extension of the Garlock Fault. Site specific geologic units and nearby faults are shown on, Figure 4-Site Vicinity Geologic Map.')
          if colorado == True:
            geomorphic_region1.add_run('The site is located with the Colorado Desert geomorphic province of California. A low-lying barren desert basin, about 245 feet below sea level in part, is dominated by the Salton Sea. The province is a depressed block between active branches of alluvium-covered San Andreas Fault with the southern extension of the Mojave Desert on the east. It is characterized by ancient beach lines and silt deposits of extinct Lake Cahuilla. Site specific geologic units and nearby faults are shown on, Figure 4-Site Vicinity Geologic Map.')
          if peninsular == True:
            geomorphic_region1.add_run('The site is located with the geomorphic Peninsualar Ranges province of California. A series of ranges is separated by northwest trending valleys, subparralel to faults branching from the San Andreas Fault. The trend of topography is similar to the Coast Ranges, but the geology is more like the Siearra Nevada, with granitic rock intruding the older metamorphic rock. The peninsular ranges extend into lower California and are bound on the east by the Colorado Desert. The island group off the southwest California coast is a part of this province. It is bordered to the north by the Transverse Ranges, and to the east by the Colorado Desert province. Site specific geologic units and nearby faults are shown on, Figure 4-Site Vicinity Geologic Map.')
          if cascade == True:
            geomorphic_region1.add_run('The site is located with the geomorphic province of California. The Cascade Ranges are a chain of volcanic cones extending through Washington and Oregon into California. It is dominated by Mt. Shasta, a glacier-mantled volcanic cone, rising 14,162 feet above sea level. The souther termination is Lassen Peak, which last erupted in the early 1900s. The Cascade Range is transected by deep canyons of the Pit River. The river flows through the range between these two major volcanic cones, after winding across interior Modoc Plateau on its way to the Sacramento River. Site specific geologic units and nearby faults are shown on, Figure 4-Site Vicinity Geologic Map.')
          if gentle == True:
            topography = 'gentle'
          if flat == True:
            topography = 'flat'
          if steep == True:
            topography = 'steep'

          table = doc.tables[3]
          site_class_insert_loc = table.cell(1,1)
          site_class_insert_loc.text = text_input_box_site_class.text
          ss_insert_loc = table.cell(2,1)
          ss_insert_loc.text = text_input_box_ss.text + "\n" + text_input_box_s1.text
          fa_insert_loc = table.cell(3,1)
          fa_insert_loc.text = text_input_box_fa.text
          fv_insert_loc = table.cell(4,1)
          fv_insert_loc.text = text_input_box_fv.text
          sms_insert_loc = table.cell(5,1)
          sms_insert_loc.text = text_input_box_sms.text
          sm1_insert_loc = table.cell(6,1)
          sm1_insert_loc.text = text_input_box_sm1.text
          sds_insert_loc = table.cell(7,1)
          sds_insert_loc.text = text_input_box_sds.text + "\n" + text_input_box_sd1.text
          

          doc.save(F'{gint_address1}_Geotechnical_Report.docx')
          gintyes = False   
          run = False                           
          
      screen.blit(text28,(200,100))
      screen.blit(text29,(200,135))

      screen.blit(textfinish,(560,210))

      screen.blit(textquit,(1100,35))
      screen.blit(textback,(50,655))

    if menu_state == "9":
      if back_button.draw(screen):
        menu_state = "8"
      if quit_button.draw(screen):
        run = False
      screen.blit(textquit,(1100,35))
      screen.blit(textback,(50,655))
      

      screen.blit(textquit,(1100,35))
      
  else:
    if quit_button.draw(screen):
      run = False
    if start_button.draw(screen):
      project_start = True
    if learning_start_button.draw(screen):
      learning_mode = True
      project_start = True

    screen.blit(textquit,(1100,35))   
    screen.blit(title_text,(360,35))
    screen.blit(main_menu_text1,(400,200))
    screen.blit(main_menu_text3,(290,325))

  #event handler
  for event in pygame.event.get():
    if event.type == pygame.KEYDOWN:
      if event.key == pygame.K_SPACE:
        project_start = True
    if event.type == pygame.QUIT:
      run = False

  pygame.display.update()

pygame.quit()